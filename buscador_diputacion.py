import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import re
import os
from docx import Document

def rastreador_diputacion_total():
    # 1. Configuración de archivos
    directorio = os.path.dirname(os.path.abspath(__file__))
    fecha_hoy_str = datetime.now().strftime("%d_%m_%Y")
    nombre_word = os.path.join(directorio, f"Diputacion_Coruna_{fecha_hoy_str}.docx")
    
    print(f"\n--- 🏛️ VIGILANCIA TOTAL: DIPUTACIÓN + RRHH ---")

    # FILTROS DE BÚSQUEDA
    # "El Quién": Buscamos específicamente a la Diputación
    terminos_entidad = [r"diputación provincial de a coruña", r"deputación da coruña"]
    # "El Qué": Convocatorias, bases o movimientos de RRHH
    terminos_bases = ["convocatoria", "bases", "proceso selectivo", "recursos humanos", "rrhh", "oferta de empleo", "recursos humans", "oferta de emprego"]

    doc = Document()
    doc.add_heading(f'Alertas Diputación A Coruña (BOE/DOG/BOP) - {datetime.now().strftime("%d/%m/%Y")}', 0)
    
    anuncios_finales = []
    hoy = datetime.now()

    # Buscamos en los últimos 7 días
    for i in range(7):
        fecha = hoy - timedelta(days=i)
        f_str = fecha.strftime("%d/%m/%Y")
        
        # Las 3 fuentes que mencionas para cubrir la entrada en vigor
        urls = {
            "BOE": fecha.strftime("https://www.boe.es/boe/dias/%Y/%m/%d/"),
            "BOP Coruña": f"https://bop.dacoruna.gal/bopportal/cambioBoletin.do?fechaInput={f_str}",
            "DOG": f"https://www.xunta.gal/diario-oficial-galicia/mostrarContenido.do?ruta=/{fecha.year}/{fecha.strftime('%Y%m%d')}/Secciones3_gl.html"
        }

        print(f"🔎 Revisando {f_str}...", end="\r")

        for fuente, url in urls.items():
            try:
                res = requests.get(url, timeout=15, headers={'User-Agent': 'Mozilla/5.0'})
                if res.status_code != 200: continue

                # Usamos lxml para procesar las tablas de los diarios oficiales
                sopa = BeautifulSoup(res.text, 'lxml')
                
                # Buscamos en todos los contenedores de texto posibles
                for item in sopa.find_all(['li', 'p', 'tr', 'td']):
                    texto = item.get_text(separator=" ").strip()
                    if len(texto) < 40: continue
                    
                    txt_min = texto.lower()

                    # Lógica de doble validación: Entidad + Palabra Clave
                    es_diputacion = any(re.search(t, txt_min) for t in terminos_entidad)
                    es_interesante = any(b in txt_min for b in terminos_bases)

                    if es_diputacion and es_interesante:
                        # Evitamos guardar dos veces la misma línea
                        if not any(a['texto'] == texto for a in anuncios_finales):
                            anuncios_finales.append({
                                'texto': texto, 
                                'fuente': fuente, 
                                'fecha': f_str, 
                                'url': url
                            })
            except:
                continue

    # 3. Generación del documento si hay resultados
    if anuncios_finales:
        for a in anuncios_finales:
            p = doc.add_paragraph()
            p.add_run(f"📌 {a['fuente']} - {a['fecha']}").bold = True
            doc.add_paragraph(a['texto'])
            doc.add_paragraph(f"🔗 {a['url']}")
            doc.add_paragraph("-" * 30)
        
        doc.save(nombre_word)
        print(f"\n✅ ¡Localizados {len(anuncios_finales)} anuncios relevantes!")
    else:
        print("\nℹ️ Sin novedades críticas de la Diputación en la última semana.")

if __name__ == "__main__":
    rastreador_diputacion_total()
